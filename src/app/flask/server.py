from flask import Flask, request, make_response, send_file, send_from_directory, jsonify
from flask_cors import CORS, cross_origin
from docx import Document
import pandas
import xlrd

def make_doc(education_level="Бакалавриат", course="Информационные технологии в дизайне", discipline="Информатика", education_start_year="2021", programm_variant="5"):

    course_discipline = "Информатика"
    if education_level == 'Бакалавриат':
        course_code = "09.03.02"
        course_profile = 'Информационные системы и технологии'
    if education_level == 'Магистратура':
        course_code = '09.04.02'
        course_profile = course
    graduate_department = "ГИС"
    developer_department = "ГИС"

    if education_level == 'Бакалавриат':
        xls_file = pandas.ExcelFile('C:/Users/Sofya/IdeaProjects/Thesis/src/app/flask/ushebn_plan/IST4.xls')
        comp_range = 4200
        comp_file = Document('C:/Users/Sofya/IdeaProjects/Thesis/src/app/flask/indikator/bac/comp_GIS.docx')
        comp_quantity = 175
        df_len = 113
    if education_level == 'Магистратура':
        comp_quantity = 75
        df_len = 80
        if course == 'Информационные технологии в дизайне':
            xls_file = pandas.ExcelFile('C:/Users/Sofya/IdeaProjects/Thesis/src/app/flask/ushebn_plan/ITD.xls')
            comp_range = 3200
            comp_file = Document('C:/Users/Sofya/IdeaProjects/Thesis/src/app/flask/indikator/mag/comp_ITD.docx')
        if course == 'Информационная поддержка жизненного цикла изделий и инфраструктуры':
            xls_file = pandas.ExcelFile('C:/Users/Sofya/IdeaProjects/Thesis/src/app/flask/ushebn_plan/IPI.xls')
            comp_range = 3000
            comp_file = Document('C:/Users/Sofya/IdeaProjects/Thesis/src/app/flask/indikator/mag/comp_IPI.docx')

    df = pandas.read_excel(xls_file, 'План')
    doc_file = Document('C:/Users/Sofya/IdeaProjects/Thesis/src/app/flask/shablon.docx')

###   Рабочая программа дисциплины
    doc_file.paragraphs[17].text = discipline

###   Рабочая программа дисциплини 2
    doc_file.paragraphs[24].text = discipline 
###   База

    doc_file.paragraphs[26].text += '   ' + course_code + ' ' + course_profile
    doc_file.paragraphs[28].text += '   ' + course
    doc_file.paragraphs[33].text += '   ' + education_start_year
    doc_file.paragraphs[35].text += '   ' + graduate_department
    doc_file.paragraphs[37].text += '   ' + developer_department

###   Объем дисциплины
    for i in range(df.shape[0]):
        if df.iat[i,3] == discipline:
            discipline_volume_amount = df.iat[i,14]
    doc_file.paragraphs[39].text += discipline_volume_amount

###   Промежуточная аттестация
    courses_index = [4,8,13,17,22,26,31,35]
    dr = pandas.read_excel(xls_file, 'Диаграмма курсов')
    attestation = []

    for i in range(8):
        for j in range(350):
            volume = str(dr.iat[j,courses_index[i]])
            if discipline in volume:
                if 'За' in volume and 'ЗаО' not in volume:
                    attestation.append('Сем ' + str(i+1) + ' - За')
                    break
                if 'ЗаО' in volume:
                    attestation.append('Сем ' + str(i+1) + ' - ЗаО')
                    break
                if 'Экз' in volume:
                    attestation.append('Сем ' + str(i+1) + ' - Экз')
                    break
    attestation_string = ''
    for i in attestation:
        attestation_string += '   ' + i
    doc_file.paragraphs[41].text += '   ' + attestation_string

    if programm_variant == "Для дисциплин философия, история, иностранный язык, безопасность жизнедеятельности, физическая культура и спорт":
        text_variant = '''2.1. Учебная дисциплина  __________________________________ включена в обязательный перечень дисциплин в рамках базовой части Блока 1, установленного ФГОС ВО, и является обязательной для всех профилей направления подготовки _____________________________ 
    Код направления подготовки
Дисциплина базируется на следующих дисциплинах: … в объёме курса средней школы / программы бакалавриата./ Предшествующими курсами, на которых непосредственно базируется дисциплина «____________________________» являются …, …, …, … .
Дисциплина «____________________________» является основополагающей для изучения следующих дисциплин: …, …, …, ….'''
    if programm_variant == 'Для дисциплин относящихся к базовой части образовательной программы':
        text_variant = '''Учебная дисциплина (модуль) ______________________ включена в обязательный перечень дисциплин обязательной части образовательной программы вне зависимости от ее направленности (профиля). Дисциплина реализуется в соответствии с требованиями ФГОС, ОП ВО и УП, по направлению  подготовки _______________________
                                                                                              Код направления подготовки
Дисциплина базируется на следующих дисциплинах: … в объёме курса средней школы / программы бакалавриата./ Предшествующими курсами, на которых непосредственно базируется дисциплина «____________________________» являются …, …, …, … .
Дисциплина «____________________________» является основополагающей для изучения следующих дисциплин: …, …, …, ….'''
    if programm_variant == 'Для дисциплин вариативной части, определяющие направленность или специализацию ОП':
        text_variant = '''Учебная дисциплина (модуль) ________________________ включена в перечень дисциплин вариативной части (формируемой участниками образовательных отношений), определяющий направленность ОП. Дисциплина реализуется в соответствии с требованиями ФГОС, ОП ВО и УП.                                                                                                     
Дисциплина базируется на следующих дисциплинах: … в объёме курса средней школы / программы бакалавриата./ Предшествующими курсами, на которых непосредственно базируется дисциплина «____________________________» являются …, …, …, … .
Дисциплина «____________________________» является основополагающей для изучения следующих дисциплин: …, …, …, ….'''
    if programm_variant == 'Для дисциплин вариативной части, углубляющих  формирование компетенций по направленности или специализации ОП':
        text_variant = '''Учебная дисциплина (модуль) ________________________ включена в перечень, вариативной части дисциплин (формируемой участниками образовательных отношений) по выбору (запросу студентов), направленный на углубление уровня освоения компетенций.  Дисциплина реализуется в соответствии с требованиями ФГОС, ОП ВО и УП. 
Дисциплина базируется на следующих дисциплинах: … в объёме курса средней школы / программы бакалавриата./ Предшествующими курсами, на которых непосредственно базируется дисциплина «____________________________» являются …, …, …, … .
Дисциплина «____________________________» является основополагающей для изучения следующих дисциплин: …, …, …, ….'''
    if programm_variant == 'Для факультативных дисциплин':
        text_variant = '''Учебная дисциплина (модуль) ________________________ включена в перечень факультативных дисциплин.
Дисциплина базируется на следующих дисциплинах: … в объёме курса средней школы / программы бакалавриата./ Предшествующими курсами, на которых непосредственно базируется дисциплина «____________________________» являются …, …, …, … .
Дисциплина «____________________________» является основополагающей для изучения следующих дисциплин: …, …, …, ….'''
    
    doc_file.paragraphs[72].text = text_variant
 
    original_string = doc_file.tables[0].cell(0,0).text
    split_strings = original_string.split()
    split_strings.insert(19, course_code + ' ' + course)
    final_string = ' '.join(split_strings)
    doc_file.tables[0].cell(0,0).text = final_string


    ###   Компетенции 
    dk = pandas.read_excel(xls_file, 'Компетенции')

    comp_discipline = []
    discipline_comp = []
    discipline_sem = []
    for i in range(comp_range):
        volume = str(dk.iat[i,6])
        if discipline in volume:
            volume_comp = str(dk.iat[(i//201*201-1),3])
            if volume_comp == 'nan': 
                volume_comp = 'ОПК-1'
            for j in range((i//201*201),((i//201)+1)*201-1):
                volume_discipline = str(dk.iat[j,6])
                if volume_discipline != "nan":
                    comp_discipline.append(volume_discipline)
                    discipline_comp.append(volume_comp)

    for i in range(len(comp_discipline)):
        discipline_sem.append([])
        for j in range(8):
            for k in range(350):
                volume = str(dr.iat[k,courses_index[j]])
                if comp_discipline[i] in volume:
                    discipline_sem[i].append(j+1)

    for i in range(len(discipline_comp)):
        doc_file.tables[1].add_row()
        doc_file.tables[1].cell(i+2,0).text = comp_discipline[i] + '. ' + discipline_comp[i]
        for j in range(len(discipline_sem[i])):
            doc_file.tables[1].cell(i+2,discipline_sem[i][j]).text = 'V'

    ###   Компетенции 2

    comps = list(set(discipline_comp))
    comps.sort()
    
    table_iterator = 2
    for i in range(len(comps)):
        for j in range(comp_quantity):
            if comps[i] in comp_file.tables[0].cell(j,0).text:
                doc_file.tables[2].add_row()
                doc_file.tables[6].add_row()
                doc_file.tables[2].cell(table_iterator,0).text = comp_file.tables[0].cell(j,0).text
                doc_file.tables[2].cell(table_iterator,1).text = comp_file.tables[0].cell(j,1).text
                doc_file.tables[6].cell(table_iterator,0).text = comp_file.tables[0].cell(j,0).text
                doc_file.tables[6].cell(table_iterator,1).text = comp_file.tables[0].cell(j,1).text

                opk_string = comp_file.tables[0].cell(j,0).text
                split_opk = opk_string.split()
                for opks in split_opk:
                    if 'ОПК' in opks or 'УК' in opks:
                        opkas = opks

                iopk_string = comp_file.tables[0].cell(j,1).text
                split_iopk = iopk_string.split()
                iopkas = []
                for iopks in split_iopk:
                    if 'ИОПК' in iopks or 'ИУК' in iopks:
                        iopkas.append(iopks)
                doc_file.tables[4].add_row() 
                doc_file.tables[4].cell(table_iterator+16,0).text += opkas
                for iopks2 in iopkas:
                    doc_file.tables[4].cell(table_iterator+16,0).text += '\n' + iopks2
                table_iterator += 1
                break
    lec_summ = 0
    lec_all = []
    pr_summ = 0
    pr_all = []
    lab_summ = 0
    lab_all = []
    sr_summ = 0
    sr_all = []
    for i in range(df.shape[0]):
        if df.iat[i,3] == discipline:
            doc_file.tables[3].cell(4,1).text = df.iat[i,14]
            doc_file.tables[3].cell(6,1).text = df.iat[i,14]
            for j in range(28,df_len):
                if 'Лек' in str(df.iat[3,j]):
                    if type(df.iat[i,j]) != type(float('nan')):
                        lec_all.append(int(df.iat[i,j]))
                if 'Лаб' in str(df.iat[3,j]):
                    if type(df.iat[i,j]) != type(float('nan')):
                        lab_all.append(int(df.iat[i,j]))
                if 'Пр' in str(df.iat[3,j]):
                    if type(df.iat[i,j]) != type(float('nan')):
                        pr_all.append(int(df.iat[i,j]))
                if 'СР' in str(df.iat[3,j]):
                    if type(df.iat[i,j]) != type(float('nan')):
                        sr_all.append(int(df.iat[i,j]))

            for j in range(len(lec_all)):
                lec_summ += lec_all[j]
            for j in range(len(lab_all)):
                lab_summ += lab_all[j]
            for j in range(len(pr_all)):
                pr_summ += pr_all[j]
            for j in range(len(sr_all)):
                sr_summ += sr_all[j]

            if lec_summ != 0:
                doc_file.tables[3].cell(7,1).text = str(lec_summ)
            if pr_summ != 0:
                doc_file.tables[3].cell(8,1).text = str(pr_summ)
            if lab_summ != 0:
                doc_file.tables[3].cell(9,1).text = str(lab_summ)
            if sr_summ != 0:
                doc_file.tables[3].cell(14,1).text = str(sr_summ)
            break

    doc_file.save('C:/Users/Sofya/IdeaProjects/Thesis/src/assets/RPD.docx')

###   Конец функции обработки


app = Flask(__name__)
CORS(app)
app.config["DOC"]="/home/cyger/sonya_diplom/fed_backend_flask/"

@app.route("/",methods = ['POST', 'GET'])
def hello():

    request_json = request.get_json(force=True)
    education_level = request_json['urovenObraz']
    course = request_json['profil']
    discipline = request_json['disciplina']
    education_start_year = request_json['date']
    programm_variant = request_json['mestoVStructure']
    print("Start making doc")
    make_doc(education_level, course, discipline, education_start_year, programm_variant)
    print("Finish making doc")

    return jsonify(response_status='OK')


    #return send_file("/home/cyger/sonya_diplom/fed_backend_flask/test_doc.docx", as_attachment=True)

if __name__ == "__main__":


    app.run(host='0.0.0.0', port=8080)
    '''
    education_level = 'Бакалавриат'
    course = 'Информационные технологии в дизайне'

    discipline = 'Информационные технологии'

    education_start_year = '2021'

    programm_variant = 'Для дисциплин философия, история, иностранный язык, безопасность жизнедеятельности, физическая культура и спорт'

    make_doc(education_level, course, discipline, education_start_year, programm_variant)
    '''
